"""
成书编译：把已采纳的案例按"前言 + 第一章思政元素与知识图谱总览 + 按维度分章 + 附录一/二"
的固定版式拼装成一份 Word 文档。
"""
import io
from collections import defaultdict

import docx
from docx.shared import Inches

from book_front_matter import (
    PREFACE_SECTIONS, SIZHENG_ELEMENT_DEFINITIONS, GRAPH_USAGE_NOTE, CHAPTER_SUBTITLES,
)
from db import Case, CaseKnowledgeMapping, DIMENSIONS
from doc_writer import write_case_section, set_default_font
from knowledge_graph import build_graph, render_graph_png
from mermaid_tree import build_applicable_courses_mermaid
from mermaid_render import render_mermaid_batch

# 对应 README 里"五维度固定值，对应大纲第二至六章"的映射关系；"第一章"是前面
# 新增的"课程思政元素与知识图谱总览"，这里从"第二章"起刚好接上，不用挪号
CHAPTER_TITLES = {
    "政治认同": "第二章 政治认同",
    "家国情怀": "第三章 家国情怀",
    "文化素养": "第四章 文化素养",
    "宪法法治意识": "第五章 宪法法治意识",
    "道德修养": "第六章 道德修养",
}


def _write_case_overview_table(doc, cases: list):
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = "案例编号", "标题", "所属维度", "审核状态"
    for c in cases:
        cells = table.add_row().cells
        cells[0].text = c.case_code or ""
        cells[1].text = c.title or ""
        cells[2].text = c.dimension or ""
        cells[3].text = c.status or ""


def _write_knowledge_mapping_table(doc, db, cases: list):
    case_ids = [c.id for c in cases]
    mappings = (
        db.query(CaseKnowledgeMapping)
        .filter(CaseKnowledgeMapping.case_id.in_(case_ids), CaseKnowledgeMapping.status == "已采纳")
        .all()
    )
    if not mappings:
        doc.add_paragraph("（暂无已采纳的知识点关联，可在「知识点匹配」标签页运行匹配并采纳后重新导出）")
        return

    case_by_id = {c.id: c for c in cases}
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = "案例编号", "关联课程/章节", "知识点描述", "融入方式建议"
    for m in mappings:
        case = case_by_id.get(m.case_id)
        kp = m.knowledge_point
        cells = table.add_row().cells
        cells[0].text = case.case_code if case else str(m.case_id)
        cells[1].text = f"{kp.course_name}　{kp.chapter or ''}" if kp else ""
        cells[2].text = kp.description if kp else ""
        cells[3].text = m.suggestion_text or ""


def build_book_docx(db, status_filter: str = "已采纳") -> bytes:
    doc = docx.Document()
    set_default_font(doc)
    doc.add_heading("讲好数字中国故事：人工智能类课程思政案例集", level=0)

    doc.add_heading("前言", level=1)
    for title, body in PREFACE_SECTIONS:
        doc.add_heading(title, level=2)
        doc.add_paragraph(body)
    doc.add_page_break()

    doc.add_heading("第一章 课程思政元素与知识图谱总览", level=1)
    doc.add_heading("(一) 思政元素", level=2)
    for title, body in SIZHENG_ELEMENT_DEFINITIONS:
        doc.add_paragraph(f"{title}：{body}")
    doc.add_heading("(二) 知识图谱", level=2)
    graph = build_graph(db)
    png_bytes = render_graph_png(graph)
    doc.add_picture(io.BytesIO(png_bytes), width=Inches(6.3))
    doc.add_heading("(三) 总知识图谱使用说明", level=2)
    doc.add_paragraph(GRAPH_USAGE_NOTE)
    doc.add_page_break()

    cases = (
        db.query(Case)
        .filter(Case.status == status_filter)
        .order_by(Case.case_code)
        .all()
    )
    by_dimension = defaultdict(list)
    for c in cases:
        by_dimension[c.dimension].append(c)

    # 树状图渲染（起无头浏览器）比较慢，全书可能有几十个案例，这里对所有案例只批量渲染
    # 一次，不要在下面的分章节循环里逐个案例单独渲染
    case_dicts = {c.id: c.to_dict() for c in cases}
    tree_pngs = dict(zip(
        case_dicts.keys(),
        render_mermaid_batch([build_applicable_courses_mermaid(case_dicts[cid]) for cid in case_dicts]),
    ))

    for dim in DIMENSIONS:
        chapter_title = CHAPTER_TITLES.get(dim, dim)
        subtitle = CHAPTER_SUBTITLES.get(dim)
        doc.add_heading(f"{chapter_title}：{subtitle}" if subtitle else chapter_title, level=1)
        dim_cases = by_dimension.get(dim, [])
        if not dim_cases:
            doc.add_paragraph(f"（{dim}维度暂无「{status_filter}」状态的案例）")
            continue
        for c in dim_cases:
            write_case_section(doc, case_dicts[c.id], heading_level=2, course_tree_png=tree_pngs.get(c.id))

    doc.add_heading("附录一 案例总览表", level=1)
    _write_case_overview_table(doc, cases)
    doc.add_page_break()

    doc.add_heading("附录二 课程思政案例与人工智能相关课程关联信息总表", level=1)
    _write_knowledge_mapping_table(doc, db, cases)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
