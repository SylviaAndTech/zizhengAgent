"""
把一个案例的七段式内容写入 python-docx 的 Document 对象。
供「勾选案例导出Word」和「成书编译」两处共用，避免两套导出各写一份格式。
"""
import re

from docx.oxml.ns import qn

# 案例正文里的句级引用标注，如 [素材1]、[素材2:关键词片段]（第二种带定位短语，供前端hover预览用）；
# 导出成Word给人看的最终稿时，这些标注是给内部溯源用的，不应该出现在正式文档里，统一去掉
_CITATION_RE = re.compile(r"\[素材\d+(?::[^\]]*)?\]")

DEFAULT_FONT = "宋体"
_STYLED_NAMES = ("Normal", "Title", "Heading 1", "Heading 2", "Heading 3", "Heading 4", "Heading 5", "Heading 6")


def strip_citations(text: str) -> str:
    return _CITATION_RE.sub("", text or "")


def set_default_font(doc, font_name: str = DEFAULT_FONT):
    """python-docx设置字体要中英文分开设：font.name只管西文(ascii/hAnsi)，中文走的是
    w:eastAsia这个单独的属性，两个都不设的话Word里中文还是显示成默认字体，看不出变化"""
    for style_name in _STYLED_NAMES:
        try:
            style = doc.styles[style_name]
        except KeyError:
            continue
        style.font.name = font_name
        rpr = style.element.get_or_add_rPr()
        rpr.rFonts.set(qn("w:eastAsia"), font_name)


def write_case_section(doc, case_dict: dict, heading_level: int = 1):
    """case_dict: Case.to_dict() 的结果"""
    d = case_dict
    doc.add_heading(f"案例{d['case_code']}：{d['title'] or ''}", level=heading_level)
    doc.add_paragraph(f"所属维度：{d['dimension'] or ''}　审核状态：{d['status']}")

    sub = heading_level + 1

    doc.add_heading("一、完整案例", level=sub)
    doc.add_paragraph(strip_citations(d["full_narrative"]))

    doc.add_heading("二、案例教学目标", level=sub)
    for k, v in (d["teaching_objectives"] or {}).items():
        doc.add_paragraph(f"{k}：{v}")

    doc.add_heading("三、课程思政元素", level=sub)
    for k, v in (d["sizheng_elements"] or {}).items():
        if k == "六维度体现" and isinstance(v, dict):
            doc.add_paragraph("六维度体现：")
            for dim, text in v.items():
                doc.add_paragraph(f"　　{dim}：{text}")
        else:
            doc.add_paragraph(f"{k}：{v}")

    doc.add_heading("四、适用课程举例", level=sub)
    courses = d["applicable_courses"] or []
    if courses:
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text = "课程名称", "适用章节", "融入方式建议"
        for row in courses:
            cells = table.add_row().cells
            cells[0].text = str(row.get("课程名称", ""))
            cells[1].text = str(row.get("适用章节", ""))
            cells[2].text = str(row.get("融入方式建议", ""))

    doc.add_heading("五、教学设计", level=sub)
    for k, v in (d["teaching_design"] or {}).items():
        doc.add_paragraph(f"{k}：{v}")

    doc.add_heading("六、课程评价与成效", level=sub)
    for k, v in (d["evaluation"] or {}).items():
        doc.add_paragraph(f"{k}：{v}")

    doc.add_heading("七、延伸阅读", level=sub)
    for r in (d["further_reading"] or []):
        doc.add_paragraph(f"[{r.get('type', '')}] {r.get('title', '')} {r.get('url', '')}")

    doc.add_page_break()
